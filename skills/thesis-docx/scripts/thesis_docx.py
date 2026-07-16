#!/usr/bin/env python3
"""Build or restyle academic DOCX files from a reusable format spec."""

from __future__ import annotations

import argparse
from io import BytesIO
from copy import deepcopy
import json
import math
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from typing import Any, Iterator
from xml.etree import ElementTree as ET

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run


W_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NAMESPACE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
A4_WIDTH_CM = 21.0
A4_HEIGHT_CM = 29.7
CM_TO_EMU = 360000
HNU_PROFILE = "course-paper"
CITEKEY_PATTERN = re.compile(r"@([A-Za-z0-9_:\-]+)")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    subparsers = parser.add_subparsers(dest="command", required=True)

    build_parser = subparsers.add_parser("build", help="Build a thesis DOCX.")
    build_parser.add_argument("--spec", required=True, type=Path, help="Path to a format spec JSON file")
    build_parser.add_argument("--output", required=True, type=Path, help="Output DOCX path")
    build_parser.add_argument("--markdown", type=Path, help="Markdown source for Pandoc build")
    build_parser.add_argument("--bib", type=Path, help="BibTeX file for citations")
    build_parser.add_argument("--csl", type=Path, help="Optional CSL file")
    build_parser.add_argument("--metadata", type=Path, help="Metadata JSON for blank skeleton builds")
    build_parser.add_argument("--no-toc", action="store_true", help="Do not ask Pandoc to inject an automatic TOC")

    restyle_parser = subparsers.add_parser("restyle", help="Restyle an existing thesis DOCX.")
    restyle_parser.add_argument("--input", required=True, type=Path, help="Existing DOCX path")
    restyle_parser.add_argument("--spec", required=True, type=Path, help="Path to a format spec JSON file")
    restyle_parser.add_argument("--output", required=True, type=Path, help="Output DOCX path")

    return parser.parse_args()


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def ensure_parent_dir(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def skill_root() -> Path:
    return Path(__file__).resolve().parents[1]


def default_metadata_path() -> Path:
    return skill_root() / "assets" / "templates" / "metadata.json"


def load_metadata_with_fallback(path: Path | None) -> dict[str, Any]:
    source = path or default_metadata_path()
    return load_json(source)


def extract_citekeys(markdown_path: Path) -> list[str]:
    text = markdown_path.read_text(encoding="utf-8")
    keys: list[str] = []
    seen: set[str] = set()
    for match in CITEKEY_PATTERN.finditer(text):
        key = match.group(1)
        if key not in seen:
            keys.append(key)
            seen.add(key)
    return keys


def create_placeholder_bib(markdown_path: Path, temp_dir: Path) -> Path | None:
    citekeys = extract_citekeys(markdown_path)
    if not citekeys:
        return None

    placeholder_path = temp_dir / "placeholder-references.bib"
    entries = []
    for key in citekeys:
        entries.append(
            "\n".join(
                [
                    f"@misc{{{key},",
                    f"  title = {{占位参考文献：{key}}},",
                    "  author = {待补充},",
                    "  year = {2026}",
                    "}",
                ]
            )
        )
    placeholder_path.write_text("\n\n".join(entries) + "\n", encoding="utf-8")
    return placeholder_path


def resolve_spec_relative_path(spec_path: Path, candidate: str | None) -> Path | None:
    if not candidate:
        return None
    candidate_path = Path(candidate)
    if candidate_path.is_absolute():
        return candidate_path
    return (spec_path.parent / candidate_path).resolve()


def build_from_markdown(
    markdown: Path,
    output: Path,
    bib: Path | None,
    csl: Path | None,
    spec_path: Path,
    spec: dict[str, Any],
    include_toc: bool = True,
) -> None:
    pandoc_path = shutil.which("pandoc")
    if not pandoc_path:
        raise SystemExit("Pandoc is required for --markdown builds but was not found in PATH.")

    command = [pandoc_path, str(markdown), "--number-sections", "-o", str(output)]
    command.extend(["--resource-path", str(markdown.parent)])
    if include_toc:
        command.append("--toc")
    citation_spec = spec.get("citations", {})
    if bib:
        command.extend(["--citeproc", "--bibliography", str(bib)])
        effective_csl = csl or resolve_spec_relative_path(spec_path, citation_spec.get("csl"))
        if effective_csl:
            command.extend(["--csl", str(effective_csl)])
        if citation_spec.get("link_citations", False):
            command.extend(["-M", "link-citations=true"])
        if citation_spec.get("reference_section_title"):
            command.extend(["-M", f"reference-section-title={citation_spec['reference_section_title']}"])
    subprocess.run(command, check=True)


def set_document_core_properties(document: DocumentObject, metadata: dict[str, Any]) -> None:
    props = document.core_properties
    title = metadata.get("title") or metadata.get("title_cn")
    author = metadata.get("author")
    if title:
        props.title = title
    if author:
        props.author = author


def set_update_fields_on_open(docx_path: Path) -> None:
    tag = f"{{{W_NAMESPACE}}}updateFields"
    attr = f"{{{W_NAMESPACE}}}val"

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        with zipfile.ZipFile(docx_path, "r") as source_zip:
            source_zip.extractall(temp_path)

        settings_path = temp_path / "word" / "settings.xml"
        tree = ET.parse(settings_path)
        root = tree.getroot()
        update_node = root.find(tag)
        if update_node is None:
            update_node = ET.SubElement(root, tag)
        update_node.set(attr, "true")
        tree.write(settings_path, encoding="utf-8", xml_declaration=True)

        rebuilt_path = temp_path / "rebuilt.docx"
        with zipfile.ZipFile(rebuilt_path, "w", compression=zipfile.ZIP_DEFLATED) as target_zip:
            for file_path in sorted(temp_path.rglob("*")):
                if file_path.is_dir() or file_path == rebuilt_path:
                    continue
                target_zip.write(file_path, file_path.relative_to(temp_path))
        shutil.move(rebuilt_path, docx_path)


def clear_update_fields_on_open(docx_path: Path) -> None:
    tag = f"{{{W_NAMESPACE}}}updateFields"

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        with zipfile.ZipFile(docx_path, "r") as source_zip:
            source_zip.extractall(temp_path)

        settings_path = temp_path / "word" / "settings.xml"
        tree = ET.parse(settings_path)
        root = tree.getroot()
        update_node = root.find(tag)
        if update_node is not None:
            root.remove(update_node)
        tree.write(settings_path, encoding="utf-8", xml_declaration=True)

        rebuilt_path = temp_path / "rebuilt.docx"
        with zipfile.ZipFile(rebuilt_path, "w", compression=zipfile.ZIP_DEFLATED) as target_zip:
            for file_path in sorted(temp_path.rglob("*")):
                if file_path.is_dir() or file_path == rebuilt_path:
                    continue
                target_zip.write(file_path, file_path.relative_to(temp_path))
        shutil.move(rebuilt_path, docx_path)


def ensure_style_xml(styles_root, style_id: str, style_name: str):
    style_tag = qn("w:style")
    name_tag = qn("w:name")
    based_on_tag = qn("w:basedOn")
    for style in styles_root.findall(style_tag):
        if style.get(qn("w:styleId")) == style_id:
            break
    else:
        style = ET.SubElement(styles_root, style_tag)
        style.set(qn("w:type"), "paragraph")
        style.set(qn("w:styleId"), style_id)

    style.attrib.pop(qn("w:customStyle"), None)

    name = style.find(name_tag)
    if name is None:
        name = ET.SubElement(style, name_tag)
    name.set(qn("w:val"), style_name)

    based_on = style.find(based_on_tag)
    if based_on is None:
        based_on = ET.SubElement(style, based_on_tag)
    based_on.set(qn("w:val"), "Normal")
    return style


def patch_rpr_xml(rpr, font_spec: dict[str, Any]) -> None:
    def find_or_add(parent, tag):
        node = parent.find(tag)
        if node is None:
            node = ET.SubElement(parent, tag)
        return node

    rfonts = find_or_add(rpr, qn("w:rFonts"))
    western = font_spec.get("western", "Times New Roman")
    east_asia = font_spec.get("east_asia", western)
    rfonts.set(qn("w:ascii"), western)
    rfonts.set(qn("w:hAnsi"), western)
    rfonts.set(qn("w:cs"), western)
    rfonts.set(qn("w:eastAsia"), east_asia)

    if font_spec.get("bold"):
        find_or_add(rpr, qn("w:b"))
        find_or_add(rpr, qn("w:bCs"))
    else:
        for tag in (qn("w:b"), qn("w:bCs")):
            node = rpr.find(tag)
            if node is not None:
                rpr.remove(node)

    color = find_or_add(rpr, qn("w:color"))
    color.set(qn("w:val"), font_spec.get("color", "000000").replace("#", ""))

    size_val = str(int(round(font_spec.get("size_pt", 12) * 2)))
    sz = find_or_add(rpr, qn("w:sz"))
    sz.set(qn("w:val"), size_val)
    sz_cs = find_or_add(rpr, qn("w:szCs"))
    sz_cs.set(qn("w:val"), size_val)


def patch_ppr_xml(ppr, style_spec: dict[str, Any]) -> None:
    def find_or_add(parent, tag):
        node = parent.find(tag)
        if node is None:
            node = ET.SubElement(parent, tag)
        return node

    spacing = find_or_add(ppr, qn("w:spacing"))
    spacing.set(qn("w:before"), str(int(round(style_spec.get("space_before_pt", 0) * 20))))
    spacing.set(qn("w:after"), str(int(round(style_spec.get("space_after_pt", 0) * 20))))
    if "line_spacing" in style_spec:
        spacing.set(qn("w:line"), str(int(round(style_spec["line_spacing"] * 240))))
        spacing.set(qn("w:lineRule"), "auto")

    if "left_indent_cm" in style_spec or "first_line_indent_pt" in style_spec:
        ind = find_or_add(ppr, qn("w:ind"))
        if "left_indent_cm" in style_spec:
            ind.set(qn("w:left"), str(int(round(style_spec["left_indent_cm"] * 567))))
        if "first_line_indent_pt" in style_spec:
            ind.set(qn("w:firstLine"), str(int(round(style_spec["first_line_indent_pt"] * 20))))


def patch_builtin_toc_styles(docx_path: Path, spec: dict[str, Any]) -> None:
    styles_spec = spec.get("styles", {})
    toc1_spec = styles_spec.get("toc_entry_level1")
    toc2_spec = styles_spec.get("toc_entry_level2", toc1_spec)
    toc3_spec = styles_spec.get("toc_entry_level3", toc2_spec)
    if not toc1_spec:
        return

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        with zipfile.ZipFile(docx_path, "r") as source_zip:
            source_zip.extractall(temp_path)

        styles_path = temp_path / "word" / "styles.xml"
        tree = ET.parse(styles_path)
        root = tree.getroot()

        toc_style_specs = {
            "TOC1": ("toc 1", toc1_spec),
            "TOC2": ("toc 2", toc2_spec),
            "TOC3": ("toc 3", toc3_spec),
        }
        for style_id, (style_name, style_spec) in toc_style_specs.items():
            style = ensure_style_xml(root, style_id, style_name)
            ppr = style.find(qn("w:pPr"))
            if ppr is None:
                ppr = ET.SubElement(style, qn("w:pPr"))
            patch_ppr_xml(ppr, style_spec)
            rpr = style.find(qn("w:rPr"))
            if rpr is None:
                rpr = ET.SubElement(style, qn("w:rPr"))
            patch_rpr_xml(rpr, style_spec["font"])

        tree.write(styles_path, encoding="utf-8", xml_declaration=True)

        rebuilt_path = temp_path / "rebuilt.docx"
        with zipfile.ZipFile(rebuilt_path, "w", compression=zipfile.ZIP_DEFLATED) as target_zip:
            for file_path in sorted(temp_path.rglob("*")):
                if file_path.is_dir() or file_path == rebuilt_path:
                    continue
                target_zip.write(file_path, file_path.relative_to(temp_path))
        shutil.move(rebuilt_path, docx_path)


def get_style(document: DocumentObject, style_name: str):
    for style in document.styles:
        if style.name == style_name:
            return style
    return None


def get_or_create_paragraph_style(document: DocumentObject, style_name: str, base_style: str = "Normal"):
    style = get_style(document, style_name)
    if style is None:
        style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        base = get_style(document, base_style)
        if base is not None:
            style.base_style = base
    return style


def ensure_rpr(element) -> Any:
    rpr = element.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        element.append(rpr)
    return rpr


def ensure_ppr(element) -> Any:
    ppr = element.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        element.insert(0, ppr)
    return ppr


def ensure_rfonts(rpr) -> Any:
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    return rfonts


def set_character_spacing(rpr, spacing_pt: float) -> None:
    spacing = rpr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        rpr.append(spacing)
    spacing.set(qn("w:val"), str(int(round(spacing_pt * 20))))


def set_outline_level_on_paragraph(paragraph: Paragraph, level: int) -> None:
    ppr = paragraph._element.get_or_add_pPr()
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        ppr.append(outline)
    outline.set(qn("w:val"), str(level))


def parse_color(font_spec: dict[str, Any]) -> RGBColor:
    color = font_spec.get("color", "000000").replace("#", "")
    if len(color) != 6:
        color = "000000"
    return RGBColor.from_string(color)


def apply_font_to_style(style, font_spec: dict[str, Any]) -> None:
    western = font_spec.get("western")
    east_asia = font_spec.get("east_asia", western)
    size_pt = font_spec.get("size_pt")
    if western:
        style.font.name = western
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    style.font.color.rgb = parse_color(font_spec)
    if "bold" in font_spec:
        style.font.bold = font_spec["bold"]
    if "italic" in font_spec:
        style.font.italic = font_spec["italic"]

    if western or east_asia or "character_spacing_pt" in font_spec:
        rpr = ensure_rpr(style.element)
        rfonts = ensure_rfonts(rpr)
        if western:
            rfonts.set(qn("w:ascii"), western)
            rfonts.set(qn("w:hAnsi"), western)
            rfonts.set(qn("w:cs"), western)
        if east_asia:
            rfonts.set(qn("w:eastAsia"), east_asia)
        if "character_spacing_pt" in font_spec:
            set_character_spacing(rpr, font_spec["character_spacing_pt"])


def apply_font_to_run(run: Run, font_spec: dict[str, Any]) -> None:
    western = font_spec.get("western")
    east_asia = font_spec.get("east_asia", western)
    size_pt = font_spec.get("size_pt")
    if western:
        run.font.name = western
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run.font.color.rgb = parse_color(font_spec)
    if "bold" in font_spec:
        run.font.bold = font_spec["bold"]
    if "italic" in font_spec:
        run.font.italic = font_spec["italic"]

    if western or east_asia or "character_spacing_pt" in font_spec:
        rpr = ensure_rpr(run._element)
        rfonts = ensure_rfonts(rpr)
        if western:
            rfonts.set(qn("w:ascii"), western)
            rfonts.set(qn("w:hAnsi"), western)
            rfonts.set(qn("w:cs"), western)
        if east_asia:
            rfonts.set(qn("w:eastAsia"), east_asia)
        if "character_spacing_pt" in font_spec:
            set_character_spacing(rpr, font_spec["character_spacing_pt"])


def alignment_value(name: str | None):
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(name or "left")


def apply_paragraph_spec(paragraph: Paragraph, spec: dict[str, Any]) -> None:
    paragraph.alignment = alignment_value(spec.get("alignment"))
    fmt = paragraph.paragraph_format
    if "line_spacing" in spec:
        fmt.line_spacing = spec["line_spacing"]
    if "space_before_pt" in spec:
        fmt.space_before = Pt(spec["space_before_pt"])
    if "space_after_pt" in spec:
        fmt.space_after = Pt(spec["space_after_pt"])
    if "first_line_indent_pt" in spec:
        fmt.first_line_indent = Pt(spec["first_line_indent_pt"])
    if "first_line_indent_cm" in spec:
        fmt.first_line_indent = Cm(spec["first_line_indent_cm"])
    if "left_indent_cm" in spec:
        fmt.left_indent = Cm(spec["left_indent_cm"])
    if "right_indent_cm" in spec:
        fmt.right_indent = Cm(spec["right_indent_cm"])
    if "keep_with_next" in spec:
        fmt.keep_with_next = spec["keep_with_next"]
    if "page_break_before" in spec:
        fmt.page_break_before = spec["page_break_before"]


def apply_style_spec(style, spec: dict[str, Any]) -> None:
    apply_font_to_style(style, spec["font"])
    fmt = style.paragraph_format
    if "line_spacing" in spec:
        fmt.line_spacing = spec["line_spacing"]
    if "space_before_pt" in spec:
        fmt.space_before = Pt(spec["space_before_pt"])
    if "space_after_pt" in spec:
        fmt.space_after = Pt(spec["space_after_pt"])
    if "first_line_indent_pt" in spec:
        fmt.first_line_indent = Pt(spec["first_line_indent_pt"])
    if "first_line_indent_cm" in spec:
        fmt.first_line_indent = Cm(spec["first_line_indent_cm"])
    if "left_indent_cm" in spec:
        fmt.left_indent = Cm(spec["left_indent_cm"])
    if "right_indent_cm" in spec:
        fmt.right_indent = Cm(spec["right_indent_cm"])
    if "keep_with_next" in spec:
        fmt.keep_with_next = spec["keep_with_next"]
    if "page_break_before" in spec:
        fmt.page_break_before = spec["page_break_before"]


def configure_styles(document: DocumentObject, spec: dict[str, Any]) -> None:
    styles = spec["styles"]
    mappings = {
        "Normal": styles["body"],
        "Body Text": styles["body"],
        "First Paragraph": styles["body"],
        "Title": styles["title"],
        "Subtitle": styles.get("subtitle", styles["title"]),
        "Heading 1": styles["heading1"],
        "Heading 2": styles["heading2"],
        "Heading 3": styles["heading3"],
        "Caption": styles["caption"],
    }

    for style_name, style_spec in mappings.items():
        style = get_style(document, style_name)
        if style is not None:
            apply_style_spec(style, style_spec)

    toc1_spec = styles.get("toc_entry_level1")
    toc2_spec = styles.get("toc_entry_level2", toc1_spec)
    toc3_spec = styles.get("toc_entry_level3", toc2_spec)
    toc_style_specs = {
        "TOC 1": toc1_spec,
        "TOC 2": toc2_spec,
        "TOC 3": toc3_spec,
    }
    for style_name, style_spec in toc_style_specs.items():
        if style_spec:
            style = get_or_create_paragraph_style(document, style_name)
            style.element.attrib.pop(qn("w:customStyle"), None)
            apply_style_spec(style, style_spec)

    optional_styles = {
        "Thesis Meta": styles.get("meta"),
        "Thesis Subtitle": styles.get("subtitle"),
        "Thesis Reference Entry": styles.get("reference"),
        "Figure Caption": styles.get("caption"),
        "Table Caption": styles.get("caption"),
        "HNU Cover Major": styles.get("cover_major"),
        "HNU Cover Title CN": styles.get("cover_title_cn"),
        "HNU Cover Title EN": styles.get("cover_title_en"),
        "HNU Cover Info": styles.get("cover_info"),
        "HNU Cover Date": styles.get("cover_date"),
        "HNU Front Heading": styles.get("front_heading"),
        "HNU Abstract Title CN": styles.get("abstract_title_cn"),
        "HNU Abstract Title EN": styles.get("abstract_title_en"),
        "HNU Abstract EN Body": styles.get("abstract_body_en"),
    }
    for style_name, style_spec in optional_styles.items():
        if style_spec:
            apply_style_spec(get_or_create_paragraph_style(document, style_name), style_spec)


def configure_page_layout(document: DocumentObject, spec: dict[str, Any]) -> None:
    page = spec["page"]
    margins = page["margins_cm"]
    footer_spec = spec.get("footer", {})
    for section in document.sections:
        section.page_width = Cm(A4_WIDTH_CM)
        section.page_height = Cm(A4_HEIGHT_CM)
        section.top_margin = Cm(margins["top"])
        section.bottom_margin = Cm(margins["bottom"])
        section.left_margin = Cm(margins["left"])
        section.right_margin = Cm(margins["right"])
        section.header_distance = Cm(page.get("header_distance_cm", 1.27))
        section.footer_distance = Cm(page.get("footer_distance_cm", 1.27))
        section.different_first_page_header_footer = footer_spec.get("hide_on_first_page", False)

    cover_page = spec.get("cover_page")
    if spec.get("profile") == HNU_PROFILE and cover_page and document.sections:
        cover_margins = cover_page["margins_cm"]
        cover_section = document.sections[0]
        cover_section.top_margin = Cm(cover_margins["top"])
        cover_section.bottom_margin = Cm(cover_margins["bottom"])
        cover_section.left_margin = Cm(cover_margins["left"])
        cover_section.right_margin = Cm(cover_margins["right"])


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._element):
        paragraph._element.remove(child)


def clear_paragraph_runs(paragraph: Paragraph) -> None:
    for child in list(paragraph._element):
        if child.tag != qn("w:pPr"):
            paragraph._element.remove(child)


def clear_story(story) -> Paragraph:
    element = story._element
    for child in list(element):
        element.remove(child)
    return story.add_paragraph()


def add_field(paragraph: Paragraph, instruction: str, placeholder_text: str) -> Run:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = placeholder_text

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_separate)
    run._element.append(text)
    run._element.append(fld_end)
    return run


def set_paragraph_bottom_border(paragraph: Paragraph, border_style: str, size: int = 8) -> None:
    ppr = paragraph._element.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = pbdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pbdr.append(bottom)
    bottom.set(qn("w:val"), border_style)
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")


def set_section_page_numbering(section, number_format: str | None, start: int | None) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    if number_format:
        pg_num_type.set(qn("w:fmt"), number_format)
    if start is not None:
        pg_num_type.set(qn("w:start"), str(start))


def usable_width_cm(section) -> float:
    return (section.page_width - section.left_margin - section.right_margin) / CM_TO_EMU


def add_tab_run(paragraph: Paragraph) -> None:
    paragraph.add_run().add_tab()


def add_text_segment(paragraph: Paragraph, text: str, font_spec: dict[str, Any]) -> None:
    run = paragraph.add_run(text)
    apply_font_to_run(run, font_spec)


def add_field_segment(paragraph: Paragraph, field_kind: str, font_spec: dict[str, Any]) -> None:
    if field_kind == "heading1":
        run = add_field(paragraph, ' STYLEREF "Heading 1" ', "当前章节")
    elif field_kind == "page-upper-roman":
        run = add_field(paragraph, " PAGE \\* ROMAN ", "I")
    else:
        run = add_field(paragraph, " PAGE ", "1")
    apply_font_to_run(run, font_spec)


def configure_story_line(story, section, config: dict[str, Any]) -> None:
    paragraph = clear_story(story)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    font_spec = config["font"]

    center_stop = usable_width_cm(section) / 2
    right_stop = usable_width_cm(section)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(center_stop), WD_TAB_ALIGNMENT.CENTER)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(right_stop), WD_TAB_ALIGNMENT.RIGHT)

    if config.get("left_text"):
        add_text_segment(paragraph, config["left_text"], font_spec)
    if config.get("left_field"):
        add_field_segment(paragraph, config["left_field"], font_spec)
    add_tab_run(paragraph)

    if config.get("center_text"):
        add_text_segment(paragraph, config["center_text"], font_spec)
    if config.get("center_field"):
        add_field_segment(paragraph, config["center_field"], font_spec)
    add_tab_run(paragraph)

    if config.get("right_text"):
        add_text_segment(paragraph, config["right_text"], font_spec)
    if config.get("right_field"):
        add_field_segment(paragraph, config["right_field"], font_spec)

    if config.get("bottom_border"):
        set_paragraph_bottom_border(paragraph, config["bottom_border"], config.get("border_size", 8))


def clear_section_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    clear_story(section.header)
    clear_story(section.footer)


def apply_hnu_headers_and_footers(document: DocumentObject, spec: dict[str, Any]) -> None:
    headers = spec.get("headers", {})
    footers = spec.get("footers", {})
    sections = list(document.sections)

    for section in sections:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

    if len(sections) >= 3:
        clear_section_header_footer(sections[0])
        configure_story_line(sections[1].header, sections[1], headers["front"])
        configure_story_line(sections[1].footer, sections[1], footers["front"])
        set_section_page_numbering(sections[1], footers["front"].get("page_number_format"), 1)

        configure_story_line(sections[2].header, sections[2], headers["main"])
        configure_story_line(sections[2].footer, sections[2], footers["main"])
        set_section_page_numbering(sections[2], footers["main"].get("page_number_format"), 1)

        for section in sections[3:]:
            configure_story_line(section.header, section, headers["main"])
            configure_story_line(section.footer, section, footers["main"])
    else:
        configure_story_line(sections[0].header, sections[0], headers["main"])
        configure_story_line(sections[0].footer, sections[0], footers["main"])
        set_section_page_numbering(sections[0], footers["main"].get("page_number_format"), 1)


def add_page_numbers(document: DocumentObject, spec: dict[str, Any]) -> None:
    footer_spec = spec["footer"]
    for index, section in enumerate(document.sections):
        section.footer.is_linked_to_previous = False
        paragraph = clear_story(section.footer)
        paragraph.alignment = alignment_value(footer_spec.get("alignment"))
        run = add_field(paragraph, " PAGE ", "1")
        apply_font_to_run(run, footer_spec["font"])
        if index == 0:
            set_section_page_numbering(section, "decimal", 1)


def configure_section_headers_and_footers(document: DocumentObject, spec: dict[str, Any]) -> None:
    if spec.get("profile") == HNU_PROFILE:
        apply_hnu_headers_and_footers(document, spec)
    else:
        add_page_numbers(document, spec)


def add_toc(document: DocumentObject, labels: dict[str, Any]) -> None:
    heading = document.add_paragraph(labels["toc"])
    heading.style = document.styles["Heading 1"]
    paragraph = document.add_paragraph()
    add_field(paragraph, ' TOC \\o "1-3" \\h \\z \\u ', "Right-click and update field.")


def add_generated_list_page(document: DocumentObject, title: str, sequence_name: str) -> None:
    heading = document.add_paragraph(title)
    front_heading = get_style(document, "HNU Front Heading")
    if front_heading is not None:
        heading.style = front_heading
    else:
        heading.style = document.styles["Heading 1"]
    paragraph = document.add_paragraph()
    if sequence_name == "TOC":
        add_field(paragraph, ' TOC \\o "1-3" \\h \\z \\u ', "Right-click and update field.")
    elif sequence_name == "Figure Caption":
        add_field(paragraph, ' TOC \\h \\z \\t "Figure Caption,1" ', "Right-click and update field.")
    elif sequence_name == "Table Caption":
        add_field(paragraph, ' TOC \\h \\z \\t "Table Caption,1" ', "Right-click and update field.")
    else:
        add_field(paragraph, f' TOC \\h \\z \\c "{sequence_name}" ', "Right-click and update field.")


def add_keywords_paragraph(document: DocumentObject, label: str, values: list[str], separator: str) -> None:
    if not values:
        return
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    label_run = paragraph.add_run(f"{label}：")
    label_run.bold = True
    paragraph.add_run(separator.join(values))


def has_heading(document: DocumentObject, title: str) -> bool:
    return any(paragraph.text.strip() == title for paragraph in document.paragraphs)


def ensure_reference_section(document: DocumentObject, spec: dict[str, Any]) -> None:
    reference_title = spec["labels"]["references"][0]
    if has_heading(document, reference_title):
        return
    document.add_heading(reference_title, level=1)
    paragraph = document.add_paragraph("[1] 在此补充参考文献信息。")
    ref_style = get_style(document, "Thesis Reference Entry")
    if ref_style is not None:
        paragraph.style = ref_style


def normalize_outline_entries(entries: list[Any]) -> list[dict[str, Any]]:
    normalized: list[dict[str, Any]] = []
    for entry in entries:
        if isinstance(entry, str):
            normalized.append({"title": entry, "body": "在此填写本节正文。", "children": []})
            continue
        normalized.append(
            {
                "title": entry.get("title", "未命名章节"),
                "body": entry.get("body", "在此填写本节正文。"),
                "children": normalize_outline_entries(entry.get("children", [])),
            }
        )
    return normalized


FIGURE_CAPTION_RE = re.compile(r"^图\s*([0-9]+(?:-[0-9]+)?)\s*(.+)$")
TABLE_CAPTION_RE = re.compile(r"^表\s*([0-9]+(?:-[0-9]+)?)\s*(.+)$")
LABEL_MARKER_RE = re.compile(r"^\[\[(fig|tbl):([A-Za-z0-9_.:-]+)\]\]$")
LABEL_PREFIX_RE = re.compile(r"^\[\[(fig|tbl):([A-Za-z0-9_.:-]+)\]\]\s*(.*)$")
XREF_RE = re.compile(r"\[\[ref:(fig|tbl):([A-Za-z0-9_.:-]+)\]\]")


def paragraph_text(paragraph: Paragraph) -> str:
    return "".join(run.text for run in paragraph.runs).strip()


def paragraph_has_drawing(paragraph: Paragraph) -> bool:
    return any(child.tag == qn("w:drawing") for run in paragraph.runs for child in run._element)


def detect_caption_kind(text: str) -> str | None:
    if FIGURE_CAPTION_RE.match(text):
        return "figure"
    if TABLE_CAPTION_RE.match(text):
        return "table"
    return None


def add_bookmark_to_paragraph(paragraph: Paragraph, name: str, bookmark_id: int) -> None:
    p = paragraph._p
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(bookmark_id))
    bookmark_start.set(qn("w:name"), name)

    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(bookmark_id))

    insert_pos = 1 if len(p) > 0 and p[0].tag == qn("w:pPr") else 0
    p.insert(insert_pos, bookmark_start)
    p.append(bookmark_end)


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def add_internal_hyperlink(paragraph: Paragraph, text: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "Source Han Serif CN")
    rpr.append(rfonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rpr.append(color)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    rpr.append(sz)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), "24")
    rpr.append(sz_cs)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_pageref_field(paragraph: Paragraph, anchor: str, placeholder_text: str) -> None:
    add_field(paragraph, f' PAGEREF {anchor} \\h ', placeholder_text)


def caption_short_label(text: str, kind: str) -> str:
    pattern = FIGURE_CAPTION_RE if kind == "figure" else TABLE_CAPTION_RE
    label = "图" if kind == "figure" else "表"
    match = pattern.match(text)
    if not match:
        return label
    return f"{label}{match.group(1)}"


def collect_caption_entries(document: DocumentObject) -> tuple[list[dict[str, str]], list[dict[str, str]], dict[str, dict[str, str]]]:
    figure_entries: list[dict[str, str]] = []
    table_entries: list[dict[str, str]] = []
    reference_map: dict[str, dict[str, str]] = {}
    bookmark_id = 4000
    figure_index = 0
    table_index = 0
    pending_label: tuple[str, str] | None = None

    for paragraph in list(document.paragraphs):
        text = paragraph_text(paragraph)
        prefix = LABEL_PREFIX_RE.match(text)
        if prefix and not LABEL_MARKER_RE.match(text):
            pending_label = (prefix.group(1), prefix.group(2))
            clear_paragraph_runs(paragraph)
            if prefix.group(3):
                paragraph.add_run(prefix.group(3))
            text = paragraph_text(paragraph)
        marker = LABEL_MARKER_RE.match(text)
        if marker:
            pending_label = (marker.group(1), marker.group(2))
            remove_paragraph(paragraph)
            continue
        kind = detect_caption_kind(text)
        if kind == "figure":
            figure_index += 1
            anchor = f"fig_{figure_index}"
            if pending_label and pending_label[0] == "fig":
                anchor = f"fig_{pending_label[1].replace(':', '_').replace('-', '_').replace('.', '_')}"
            add_bookmark_to_paragraph(paragraph, anchor, bookmark_id)
            bookmark_id += 1
            entry = {"text": text, "anchor": anchor}
            figure_entries.append(entry)
            if pending_label and pending_label[0] == "fig":
                reference_map[f"fig:{pending_label[1]}"] = {"anchor": anchor, "label": caption_short_label(text, kind)}
                pending_label = None
        elif kind == "table":
            table_index += 1
            anchor = f"tbl_{table_index}"
            if pending_label and pending_label[0] == "tbl":
                anchor = f"tbl_{pending_label[1].replace(':', '_').replace('-', '_').replace('.', '_')}"
            add_bookmark_to_paragraph(paragraph, anchor, bookmark_id)
            bookmark_id += 1
            entry = {"text": text, "anchor": anchor}
            table_entries.append(entry)
            if pending_label and pending_label[0] == "tbl":
                reference_map[f"tbl:{pending_label[1]}"] = {"anchor": anchor, "label": caption_short_label(text, kind)}
                pending_label = None

    return figure_entries, table_entries, reference_map


def count_caption_kinds(document: DocumentObject) -> tuple[int, int]:
    figure_entries, table_entries, _ = collect_caption_entries(document)
    return len(figure_entries), len(table_entries)


def build_text_run(base_run_element, text: str):
    new_run = deepcopy(base_run_element)
    for child in list(new_run):
        if child.tag != qn("w:rPr"):
            new_run.remove(child)
    text_node = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        text_node.set(qn("xml:space"), "preserve")
    text_node.text = text
    new_run.append(text_node)
    return new_run


def build_hyperlink_run(base_run_element, text: str, anchor: str):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.append(build_text_run(base_run_element, text))
    return hyperlink


def replace_crossrefs_in_paragraph(paragraph: Paragraph, reference_map: dict[str, dict[str, str]]) -> None:
    for run in list(paragraph.runs):
        text = run.text
        if not text or not XREF_RE.search(text):
            continue
        parent = run._element.getparent()
        if parent is None:
            continue
        insert_at = parent.index(run._element)
        fragments: list[object] = []
        last = 0
        for match in XREF_RE.finditer(text):
            if match.start() > last:
                fragments.append(build_text_run(run._element, text[last:match.start()]))
            key = f"{match.group(1)}:{match.group(2)}"
            ref = reference_map.get(key)
            if ref:
                fragments.append(build_hyperlink_run(run._element, ref["label"], ref["anchor"]))
            else:
                fragments.append(build_text_run(run._element, text[match.start():match.end()]))
            last = match.end()
        if last < len(text):
            fragments.append(build_text_run(run._element, text[last:]))
        for fragment in reversed(fragments):
            parent.insert(insert_at, fragment)
        parent.remove(run._element)


def compact_reference_leading_space(paragraph: Paragraph) -> None:
    runs = list(paragraph.runs)
    if len(runs) < 2:
        return
    first_text = runs[0].text.strip()
    if not re.match(r"^\[\d+\]$", first_text):
        return

    for run in runs[1:]:
        if run.text in {" ", "\t", " \t", "\t "}:
            run.text = ""
            continue
        if run.text.startswith(" ") or run.text.startswith("\t"):
            run.text = run.text.lstrip(" \t")
        break


def convert_figure_table_captions(document: DocumentObject, spec: dict[str, Any]) -> None:
    for paragraph in document.paragraphs:
        text = paragraph_text(paragraph)
        fig_match = FIGURE_CAPTION_RE.match(text)
        tab_match = TABLE_CAPTION_RE.match(text)
        if fig_match:
            caption_style = get_style(document, "Figure Caption") or get_style(document, "Caption")
            if caption_style is not None:
                paragraph.style = caption_style
        elif tab_match:
            caption_style = get_style(document, "Table Caption") or get_style(document, "Caption")
            if caption_style is not None:
                paragraph.style = caption_style


def chinese_section_number(index: int) -> str:
    numerals = {
        0: "零",
        1: "一",
        2: "二",
        3: "三",
        4: "四",
        5: "五",
        6: "六",
        7: "七",
        8: "八",
        9: "九",
        10: "十",
    }
    if index <= 10:
        return numerals[index]
    if index < 20:
        return "十" + numerals[index % 10]
    tens, ones = divmod(index, 10)
    result = numerals[tens] + "十"
    if ones:
        result += numerals[ones]
    return result


def hnu_numbered_sections(entries: list[Any]) -> list[dict[str, Any]]:
    normalized = normalize_outline_entries(entries)

    def apply_numbers(items: list[dict[str, Any]], level: int, path: list[int]) -> list[dict[str, Any]]:
        numbered_items: list[dict[str, Any]] = []
        for index, item in enumerate(items, start=1):
            current_path = [*path, index]
            numbered_item = {
                "title": item["title"],
                "body": item.get("body", ""),
                "children": [],
            }
            if level == 1:
                numbered_item["title"] = f"{chinese_section_number(index)}、{item['title']}"
            else:
                numbered_item["title"] = f"{'.'.join(str(part) for part in current_path)} {item['title']}"
            numbered_item["children"] = apply_numbers(item.get("children", []), level + 1, current_path)
            numbered_items.append(numbered_item)
        return numbered_items

    return apply_numbers(normalized, 1, [])


def strip_heading_prefix(text: str) -> str:
    cleaned = re.sub(r"^\s*[0-9]+(?:\.[0-9]+)*[\t .、]*", "", text)
    cleaned = re.sub(r"^\s*[一二三四五六七八九十]+、\s*", "", cleaned)
    return cleaned.strip()


def add_outline_entries(document: DocumentObject, entries: list[Any], level: int = 1) -> None:
    for entry in normalize_outline_entries(entries):
        heading_level = min(level, 3)
        document.add_heading(entry["title"], level=heading_level)
        body = entry.get("body")
        if body:
            document.add_paragraph(body, style="Normal")
        if entry["children"]:
            add_outline_entries(document, entry["children"], level + 1)


def add_title_page(document: DocumentObject, metadata: dict[str, Any], spec: dict[str, Any]) -> None:
    styles = spec["styles"]
    if metadata.get("title"):
        paragraph = document.add_paragraph(metadata["title"])
        paragraph.style = document.styles["Title"]
        apply_paragraph_spec(paragraph, styles["title"])
    if metadata.get("subtitle"):
        paragraph = document.add_paragraph(metadata["subtitle"])
        paragraph.style = document.styles["Thesis Subtitle"]
        apply_paragraph_spec(paragraph, styles["subtitle"])

    for key in ("author", "affiliation", "date"):
        if metadata.get(key):
            paragraph = document.add_paragraph(str(metadata[key]))
            paragraph.style = document.styles["Thesis Meta"]
            apply_paragraph_spec(paragraph, styles["meta"])


def add_generic_abstract_blocks(document: DocumentObject, metadata: dict[str, Any], labels: dict[str, Any]) -> None:
    if metadata.get("abstract_zh"):
        document.add_heading(labels["abstract_zh"], level=1)
        document.add_paragraph(metadata["abstract_zh"], style="Normal")
        add_keywords_paragraph(document, labels["keywords_zh"], metadata.get("keywords_zh", []), "；")

    if metadata.get("abstract_en"):
        document.add_heading(labels["abstract_en"], level=1)
        document.add_paragraph(metadata["abstract_en"], style="Normal")
        if metadata.get("keywords_en"):
            paragraph = document.add_paragraph(style="Normal")
            label_run = paragraph.add_run(f"{labels['keywords_en']}: ")
            label_run.bold = True
            paragraph.add_run("; ".join(metadata["keywords_en"]))


def build_generic_document(metadata: dict[str, Any], spec: dict[str, Any]) -> DocumentObject:
    document = Document()
    configure_styles(document, spec)
    set_document_core_properties(document, metadata)

    add_title_page(document, metadata, spec)
    document.add_page_break()
    add_toc(document, spec["labels"])
    document.add_page_break()
    add_generic_abstract_blocks(document, metadata, spec["labels"])

    reference_labels = set(spec["labels"]["references"])
    for section_title in metadata.get("sections", []):
        document.add_heading(section_title, level=1)
        if section_title in reference_labels:
            paragraph = document.add_paragraph("[1] 在此补充参考文献信息。")
            paragraph.style = document.styles["Thesis Reference Entry"]
        else:
            document.add_paragraph("在此填写本节正文。", style="Normal")

    return document


def remove_table_borders(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), "nil")


def set_table_border(table: Table, border_name: str, val: str, size: int) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    border = borders.find(qn(f"w:{border_name}"))
    if border is None:
        border = OxmlElement(f"w:{border_name}")
        borders.append(border)
    border.set(qn("w:val"), val)
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), "000000")


def set_table_column_widths(table: Table, widths_cm: list[float], total_width_cm: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(round(total_width_cm * 567))))

    tbl_grid = table._tbl.tblGrid
    for grid_col in list(tbl_grid):
        tbl_grid.remove(grid_col)
    for width_cm in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(round(width_cm * 567))))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for cell, width_cm in zip(row.cells, widths_cm):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(round(width_cm * 567))))


def is_cover_info_table(table: Table) -> bool:
    if len(table.columns) != 2 or len(table.rows) < 3:
        return False
    first_column_texts = []
    for row in table.rows:
        text = row.cells[0].text.strip()
        first_column_texts.append(text)
    return all(text.endswith("：") for text in first_column_texts if text)


def apply_three_line_table_style(table: Table, spec: dict[str, Any], total_width_cm: float) -> None:
    if is_cover_info_table(table):
        return

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_table_borders(table)
    set_table_border(table, "top", "single", 12)
    set_table_border(table, "bottom", "single", 12)

    column_count = len(table.columns)
    text_lengths: list[int] = []
    for column_index in range(column_count):
        max_len = 1
        for row in table.rows:
            text = row.cells[column_index].text.strip()
            max_len = max(max_len, len(text.replace(" ", "")))
        text_lengths.append(max_len)

    if column_count == 3 and text_lengths[1] <= max(text_lengths[0], text_lengths[2]) * 0.4:
        middle_ratio = 0.18
        side_total = text_lengths[0] + text_lengths[2]
        left_ratio = 0.41 if side_total == 0 else 0.82 * text_lengths[0] / side_total
        right_ratio = 1 - middle_ratio - left_ratio
        ratios = [left_ratio, middle_ratio, right_ratio]
    else:
        weights = [max(1.0, math.sqrt(length)) for length in text_lengths]
        total_weight = sum(weights)
        ratios = [weight / total_weight for weight in weights]
        min_ratio = 0.12 if column_count >= 4 else 0.16 if column_count == 3 else 0.2
        adjusted = [max(min_ratio, ratio) for ratio in ratios]
        adjusted_total = sum(adjusted)
        ratios = [ratio / adjusted_total for ratio in adjusted]

    column_widths_cm = [total_width_cm * ratio for ratio in ratios]
    set_table_column_widths(table, column_widths_cm, total_width_cm)

    if table.rows:
        first_row = table.rows[0]
        for cell in first_row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            bottom = tc_borders.find(qn("w:bottom"))
            if bottom is None:
                bottom = OxmlElement("w:bottom")
                tc_borders.append(bottom)
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "8")
            bottom.set(qn("w:space"), "0")
            bottom.set(qn("w:color"), "000000")

    header_spec = spec["styles"].get("table_header", spec["styles"]["body"])
    body_spec = spec["styles"].get("table_body", spec["styles"]["body"])
    for row_index, row in enumerate(table.rows):
        current_spec = header_spec if row_index == 0 else body_spec
        for cell_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                apply_paragraph_spec(paragraph, current_spec)
                if row_index == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif column_count == 3 and cell_index == 1:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif column_count == 3 and cell_index in {0, 2}:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.first_line_indent = Pt(0)
                for run in paragraph.runs:
                    apply_font_to_run(run, current_spec["font"])


def add_hnu_cover_info_table(document: DocumentObject, metadata: dict[str, Any], spec: dict[str, Any]) -> None:
    cover_fields = metadata.get("cover_fields", [])
    if not cover_fields:
        return

    table_spec = spec.get("cover_page", {}).get("info_table", {})
    label_width_cm = table_spec.get("label_width_cm", 4.2)
    value_width_cm = table_spec.get("value_width_cm", 8.0)
    row_height_cm = table_spec.get("row_height_cm", 1.25)

    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_table_borders(table)
    font_spec = spec["styles"]["cover_info"]["font"]

    for item in cover_fields:
        row = table.add_row()
        row.height = Cm(row_height_cm)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        label_cell, value_cell = row.cells
        label_cell.width = Cm(label_width_cm)
        value_cell.width = Cm(value_width_cm)
        label_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        value_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        label_paragraph = label_cell.paragraphs[0]
        cover_info_style = get_style(document, "HNU Cover Info")
        if cover_info_style is not None:
            label_paragraph.style = cover_info_style
        label_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        label_run = label_paragraph.add_run(f"{item['label']}：")
        label_run.bold = True
        apply_font_to_run(label_run, font_spec)

        value_paragraph = value_cell.paragraphs[0]
        if cover_info_style is not None:
            value_paragraph.style = cover_info_style
        value_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        value_run = value_paragraph.add_run(item.get("value", ""))
        value_run.underline = True
        apply_font_to_run(value_run, font_spec)


def add_hnu_cover(document: DocumentObject, metadata: dict[str, Any]) -> None:
    document.add_paragraph("")

    if metadata.get("title"):
        paragraph = document.add_paragraph(metadata["title"])
        paragraph.style = get_style(document, "HNU Cover Title CN")

    if metadata.get("title_en"):
        paragraph = document.add_paragraph(metadata["title_en"])
        paragraph.style = get_style(document, "HNU Cover Title EN")

    document.add_paragraph("")
    add_hnu_cover_info_table(document, metadata, spec=metadata["_spec"])

    if metadata.get("date"):
        paragraph = document.add_paragraph(metadata["date"])
        paragraph.style = get_style(document, "HNU Cover Date")


def add_hnu_abstract_page(document: DocumentObject, metadata: dict[str, Any], spec: dict[str, Any], language: str) -> None:
    labels = spec["labels"]
    if language == "zh":
        abstract_text = metadata.get("abstract_zh")
        keywords = metadata.get("keywords_zh", [])
        heading_text = labels["abstract_zh"]
        keyword_label = labels["keywords_zh"]
        separator = "；"
        body_style = "Normal"
    else:
        abstract_text = metadata.get("abstract_en")
        keywords = metadata.get("keywords_en", [])
        heading_text = labels["abstract_en"]
        keyword_label = labels["keywords_en"]
        separator = "; "
        body_style = "HNU Abstract EN Body" if get_style(document, "HNU Abstract EN Body") is not None else "Normal"

    heading = document.add_paragraph(heading_text)
    front_heading = get_style(document, "HNU Front Heading")
    if front_heading is not None:
        heading.style = front_heading
    else:
        heading.style = document.styles["Heading 1"]
    set_outline_level_on_paragraph(heading, 0)

    if abstract_text:
        document.add_paragraph(abstract_text, style=body_style)
    if keywords:
        paragraph = document.add_paragraph(style=body_style)
        label_suffix = "：" if language == "zh" else ": "
        label_run = paragraph.add_run(f"{keyword_label}{label_suffix}")
        label_run.bold = True
        paragraph.add_run(separator.join(keywords))


def add_hnu_appendices(document: DocumentObject, metadata: dict[str, Any], spec: dict[str, Any]) -> None:
    appendices = metadata.get("appendices", [])
    appendix_prefix = spec["labels"].get("appendix_prefix", "附录")
    for index, appendix in enumerate(normalize_outline_entries(appendices), start=1):
        appendix_letter = chr(64 + index)
        document.add_heading(f"{appendix_prefix} {appendix_letter} {appendix['title']}", level=1)
        if appendix.get("body"):
            document.add_paragraph(appendix["body"], style="Normal")
        if appendix["children"]:
            add_outline_entries(document, appendix["children"], level=2)


def build_hnu_document(metadata: dict[str, Any], spec: dict[str, Any]) -> DocumentObject:
    document = Document()
    metadata = {**metadata, "_spec": spec}
    configure_styles(document, spec)
    set_document_core_properties(document, metadata)

    add_hnu_cover(document, metadata)

    document.add_section(WD_SECTION.NEW_PAGE)
    add_hnu_abstract_page(document, metadata, spec, "zh")
    document.add_page_break()
    add_hnu_abstract_page(document, metadata, spec, "en")
    document.add_page_break()
    add_generated_list_page(document, spec["labels"]["toc"], "TOC")
    document.add_page_break()
    add_generated_list_page(document, spec["labels"]["list_of_figures"], "Figure Caption")
    document.add_page_break()
    add_generated_list_page(document, spec["labels"]["list_of_tables"], "Table Caption")

    document.add_section(WD_SECTION.NEW_PAGE)
    add_outline_entries(document, hnu_numbered_sections(metadata.get("sections", [])), level=1)

    references_heading = spec["labels"]["references"][0]
    document.add_heading(references_heading, level=1)
    paragraph = document.add_paragraph("[1] 在此补充参考文献信息。")
    paragraph.style = get_style(document, "Thesis Reference Entry")

    if metadata.get("acknowledgements"):
        document.add_heading(spec["labels"]["acknowledgements"], level=1)
        document.add_paragraph(metadata["acknowledgements"], style="Normal")

    add_hnu_appendices(document, metadata, spec)
    return document


def build_hnu_front_matter_document(
    metadata: dict[str, Any],
    spec: dict[str, Any],
    figure_entries: list[dict[str, str]] | None = None,
    table_entries: list[dict[str, str]] | None = None,
) -> DocumentObject:
    document = Document()
    metadata = {**metadata, "_spec": spec}
    configure_styles(document, spec)
    set_document_core_properties(document, metadata)

    add_hnu_cover(document, metadata)

    document.add_section(WD_SECTION.NEW_PAGE)
    add_hnu_abstract_page(document, metadata, spec, "zh")
    document.add_page_break()
    add_hnu_abstract_page(document, metadata, spec, "en")
    document.add_page_break()
    add_generated_list_page(document, spec["labels"]["toc"], "TOC")
    add_manual_caption_list_page(document, spec["labels"]["list_of_figures"], figure_entries or [])
    add_manual_caption_list_page(document, spec["labels"]["list_of_tables"], table_entries or [])

    document.add_section(WD_SECTION.NEW_PAGE)
    return document


def add_manual_caption_list_page(document: DocumentObject, title: str, entries: list[dict[str, str]]) -> None:
    if not entries:
        return
    document.add_page_break()
    heading = document.add_paragraph(title)
    front_heading = get_style(document, "HNU Front Heading")
    if front_heading is not None:
        heading.style = front_heading

    for entry in entries:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(15.0), WD_TAB_ALIGNMENT.RIGHT)
        toc_style = get_style(document, "TOC 1")
        if toc_style is not None:
            paragraph.style = toc_style
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.left_indent = Pt(0)
        add_internal_hyperlink(paragraph, entry["text"], entry["anchor"])
        paragraph.add_run("\t")
        add_pageref_field(paragraph, entry["anchor"], "1")


def append_document_body(source: DocumentObject, target: DocumentObject) -> None:
    rel_map: dict[str, str] = {}
    for old_rid, rel in source.part.rels.items():
        if rel.is_external:
            rel_map[old_rid] = target.part.relate_to(rel.target_ref, rel.reltype, is_external=True)
        elif "relationships/image" in rel.reltype:
            image_part = target.part.package.get_or_add_image_part(BytesIO(rel.target_part.blob))
            rel_map[old_rid] = target.part.relate_to(image_part, rel.reltype)

    target_body = target.element.body
    target_sectpr = target_body.sectPr
    for child in source.element.body:
        if child.tag == qn("w:sectPr"):
            continue
        copied = deepcopy(child)
        for element in copied.iter():
            for attr_name, attr_value in list(element.attrib.items()):
                if attr_name.startswith(f"{{{R_NAMESPACE}}}") and attr_value in rel_map:
                    element.set(attr_name, rel_map[attr_value])
        target_body.insert(target_body.index(target_sectpr), copied)


def build_blank_document(metadata: dict[str, Any], spec: dict[str, Any]) -> DocumentObject:
    if spec.get("profile") == HNU_PROFILE:
        return build_hnu_document(metadata, spec)
    return build_generic_document(metadata, spec)


def iter_block_items(parent) -> Iterator[Paragraph | Table]:
    if isinstance(parent, DocumentObject):
        parent_element = parent.element.body
    elif isinstance(parent, _Cell):
        parent_element = parent._tc
    else:
        raise TypeError(f"Unsupported parent type: {type(parent)!r}")

    for child in parent_element.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def iter_paragraphs(parent) -> Iterator[Paragraph]:
    for block in iter_block_items(parent):
        if isinstance(block, Paragraph):
            yield block
        else:
            for row in block.rows:
                for cell in row.cells:
                    yield from iter_paragraphs(cell)


def front_heading_labels(spec: dict[str, Any]) -> set[str]:
    labels = spec.get("labels", {})
    return {
        labels.get("toc", ""),
        labels.get("list_of_figures", ""),
        labels.get("list_of_tables", ""),
        labels.get("abstract_zh", ""),
        labels.get("abstract_en", ""),
    }


def paragraph_kind(
    paragraph: Paragraph,
    spec: dict[str, Any],
    reference_labels: set[str],
    front_labels: set[str],
    in_references: bool,
) -> tuple[str, bool]:
    style_name = paragraph.style.name if paragraph.style else ""
    text = paragraph.text.strip()
    profile = spec.get("profile")
    styles = spec["styles"]

    if style_name in {"TOC Heading"} and "front_heading" in styles:
        return "front_heading", in_references

    if style_name.startswith("Heading"):
        new_in_references = text in reference_labels
        if text in front_labels and "front_heading" in styles:
            return "front_heading", new_in_references
        if style_name == "Heading 1":
            return "heading1", new_in_references
        if style_name == "Heading 2":
            return "heading2", new_in_references
        return "heading3", new_in_references

    if profile == HNU_PROFILE:
        if style_name == "Title" and "cover_title_cn" in styles:
            return "cover_title_cn", in_references
        if style_name in {"Author", "Date"} and "cover_info" in styles:
            return "cover_info", in_references

    if style_name == "Title":
        return "title", in_references
    if style_name in {"Subtitle", "Thesis Subtitle"}:
        return "subtitle", in_references
    if style_name == "Thesis Meta":
        return "meta", in_references
    if style_name in {"Caption", "Figure Caption", "Table Caption"}:
        return "caption", in_references
    if style_name == "Thesis Reference Entry":
        return "reference", in_references
    if style_name.startswith("HNU Cover"):
        return {
            "HNU Cover Major": "cover_major",
            "HNU Cover Title CN": "cover_title_cn",
            "HNU Cover Title EN": "cover_title_en",
            "HNU Cover Info": "cover_info",
            "HNU Cover Date": "cover_date",
        }.get(style_name, "body"), in_references
    if style_name == "HNU Front Heading":
        return "front_heading", in_references
    if style_name == "HNU Abstract Title CN":
        return "abstract_title_cn", in_references
    if style_name == "HNU Abstract Title EN":
        return "abstract_title_en", in_references
    if style_name == "HNU Abstract EN Body":
        return "abstract_body_en", in_references
    if in_references and text:
        return "reference", in_references
    return "body", in_references


def is_hnu_special_heading(text: str, spec: dict[str, Any], reference_labels: set[str]) -> bool:
    labels = spec.get("labels", {})
    appendix_prefix = labels.get("appendix_prefix", "附录")
    acknowledgements = labels.get("acknowledgements", "")
    stripped = text.strip()
    if stripped in reference_labels:
        return True
    if acknowledgements and stripped == acknowledgements:
        return True
    if stripped.startswith(appendix_prefix):
        return True
    return False


def restyle_document(document: DocumentObject, spec: dict[str, Any]) -> None:
    configure_styles(document, spec)
    convert_figure_table_captions(document, spec)
    configure_page_layout(document, spec)
    configure_section_headers_and_footers(document, spec)

    styles = spec["styles"]
    reference_labels = set(spec["labels"]["references"])
    front_labels = front_heading_labels(spec)
    in_references = False
    h1_counter = 0
    for paragraph in iter_paragraphs(document):
        kind, in_references = paragraph_kind(paragraph, spec, reference_labels, front_labels, in_references)
        if spec.get("profile") == HNU_PROFILE and kind == "heading1":
            base_text = strip_heading_prefix(paragraph.text)
            if not is_hnu_special_heading(base_text, spec, reference_labels):
                h1_counter += 1
                desired_text = f"{chinese_section_number(h1_counter)}、{base_text}" if base_text else paragraph.text
                if paragraph.text != desired_text:
                    clear_paragraph_runs(paragraph)
                    paragraph.add_run(desired_text)
        style_spec = styles.get(kind, styles["body"])
        apply_paragraph_spec(paragraph, style_spec)
        if kind == "reference":
            ref_style = get_style(document, "Thesis Reference Entry")
            if ref_style is not None:
                paragraph.style = ref_style
            compact_reference_leading_space(paragraph)
        elif kind == "front_heading":
            front_heading = get_style(document, "HNU Front Heading")
            if front_heading is not None:
                paragraph.style = front_heading
        if paragraph_has_drawing(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.right_indent = Pt(0)
        for run in paragraph.runs:
            apply_font_to_run(run, style_spec["font"])

    table_width_cm = usable_width_cm(document.sections[-1]) * 0.94 if document.sections else 15.0
    for table in document.tables:
        apply_three_line_table_style(table, spec, table_width_cm)


def write_document(document: DocumentObject, output: Path, spec: dict[str, Any]) -> None:
    ensure_parent_dir(output)
    document.save(output)
    if spec.get("document_behavior", {}).get("update_fields_on_open", False):
        set_update_fields_on_open(output)
    else:
        clear_update_fields_on_open(output)
    patch_builtin_toc_styles(output, spec)


def handle_build(args: argparse.Namespace) -> None:
    spec = load_json(args.spec)
    metadata = load_metadata_with_fallback(args.metadata) if spec.get("profile") == HNU_PROFILE else (load_json(args.metadata) if args.metadata else {})

    if args.markdown:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_output = Path(temp_dir) / "pandoc-output.docx"
            effective_bib = args.bib or create_placeholder_bib(args.markdown, Path(temp_dir))
            include_toc = (not args.no_toc) and spec.get("profile") != HNU_PROFILE
            build_from_markdown(args.markdown, temp_output, effective_bib, args.csl, args.spec, spec, include_toc=include_toc)
            source_document = Document(str(temp_output))
            convert_figure_table_captions(source_document, spec)
            if spec.get("profile") == HNU_PROFILE:
                figure_entries, table_entries, reference_map = collect_caption_entries(source_document)
                document = build_hnu_front_matter_document(metadata, spec, figure_entries=figure_entries, table_entries=table_entries)
                append_document_body(source_document, document)
                ensure_reference_section(document, spec)
                if metadata.get("acknowledgements"):
                    document.add_heading(spec["labels"]["acknowledgements"], level=1)
                    document.add_paragraph(metadata["acknowledgements"], style="Normal")
                add_hnu_appendices(document, metadata, spec)
                for paragraph in iter_paragraphs(document):
                    replace_crossrefs_in_paragraph(paragraph, reference_map)
            else:
                document = source_document
                if metadata:
                    set_document_core_properties(document, metadata)
            restyle_document(document, spec)
            write_document(document, args.output, spec)
        return

    document = build_blank_document(metadata, spec)
    restyle_document(document, spec)
    write_document(document, args.output, spec)


def handle_restyle(args: argparse.Namespace) -> None:
    spec = load_json(args.spec)
    document = Document(str(args.input))
    restyle_document(document, spec)
    write_document(document, args.output, spec)


def main() -> None:
    args = parse_args()
    if args.command == "build":
        handle_build(args)
    elif args.command == "restyle":
        handle_restyle(args)


if __name__ == "__main__":
    main()
